import os
import uuid
import json
import re
from datetime import datetime, timedelta, date
from typing import List, Optional, Literal, Dict, Any
from sqlalchemy import or_, and_
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from zoneinfo import ZoneInfo
import ssl
from email.utils import parseaddr, formataddr
try:
    from apscheduler.schedulers.asyncio import AsyncIOScheduler
    from apscheduler.triggers.cron import CronTrigger
except Exception:
    AsyncIOScheduler = None
    CronTrigger = None
from PyPDF2 import PdfReader
import docx
import pandas as pd

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Body, Path
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, EmailStr, Field, AnyHttpUrl
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, ForeignKey, Float, func, desc, Boolean
from sqlalchemy.orm import sessionmaker, declarative_base, Session, relationship, joinedload
from passlib.context import CryptContext
import jwt

import requests
import traceback


# RAG libs
from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels
from sentence_transformers import SentenceTransformer

# Groq SDK
from groq import Groq

# ---------------- CONFIG ----------------
SECRET_KEY = os.getenv("SECRET_KEY", "dev-secret-change")
JWT_EXPIRE_MINUTES = int(os.getenv("JWT_EXPIRE_MINUTES", "240"))
DATABASE_URL = os.getenv("DATABASE_URL")

QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION = os.getenv("COLLECTION_NAME", "psych_docs")
EMBED_MODEL = os.getenv("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
TOP_K = int(os.getenv("RAG_TOP_K", "4"))
OUT_OF_CONTEXT_THRESHOLD = float(os.getenv("OUT_OF_CONTEXT_THRESHOLD", "0.22"))

# Providers
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "mistral")

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL_DEFAULT = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")
_groq_client: Optional[Groq] = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

CORS_ALLOWED_ORIGINS = os.getenv("CORS_ALLOWED_ORIGINS", "*").split(",")

# --- Email / SMTP ---
SMTP_HOST = os.getenv("SMTP_HOST", "mailhog")   # use docker service name if using docker-compose
SMTP_PORT = int(os.getenv("SMTP_PORT", "1025"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_FROM = os.getenv("SMTP_FROM", "Psych Support <no-reply@psychsupport.local>")

# "none" | "starttls" | "ssl"
SMTP_SECURE = os.getenv("SMTP_SECURE", "none").lower()
SMTP_TIMEOUT = int(os.getenv("SMTP_TIMEOUT", "20"))
# timezone for schedules
LOCAL_TZ = ZoneInfo("Asia/Kolkata")

# Enable/disable in-process scheduler (IMPORTANT when using multiple uvicorn workers)
ENABLE_EMAIL_SCHEDULER = os.getenv("ENABLE_EMAIL_SCHEDULER", "1") == "1"

# Questionnaire rules
MIN_QUESTIONS = 7          # minimum to reach before offering score
CONTINUE_STEP = 3            # increment each time user says "continue"

# ---------------- DB ----------------
engine = create_engine(
    DATABASE_URL,
    connect_args={"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {},
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# filepath: f:\Interview\llm-challenge\backend\main.py

class Appointment(Base):
    __tablename__ = "appointments"

    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    counsellor_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    start_time = Column(DateTime, nullable=False)
    end_time = Column(DateTime, nullable=False)
    duration_minutes = Column(Integer, nullable=False)
    status = Column(String, default="pending")  # pending / accepted / rejected / closed

    # ✅ New relationships
    user = relationship("User", foreign_keys=[user_id], backref="appointments_as_user")
    counsellor = relationship("User", foreign_keys=[counsellor_id], backref="appointments_as_counsellor")

    # Other fields
    report = Column(Text, nullable=True)
    meeting_link = Column(String(512), nullable=True)
    reminder_24_sent = Column(Boolean, default=False)
    reminder_12_sent = Column(Boolean, default=False)

class Video(Base):
    __tablename__ = "videos"
    id = Column(String(36), primary_key=True, index=True)
    title = Column(String(255), nullable=False)
    url = Column(String(512), nullable=False)
    platform = Column(String(32), default="youtube")
    description = Column(Text)
    tags = Column(String(255))
    is_public = Column(Integer, default=1)
    added_by = Column(String(255), default="admin")
    access = Column(String(32), default="all")   # all | users | counsellor | admin | user:<id>
    created_at = Column(DateTime, default=datetime.utcnow)
    



class AccessRequest(Base):
    __tablename__ = "access_requests"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)          # target user
    counsellor_id = Column(Integer, ForeignKey("users.id"), index=True)    # requester
    request_type = Column(String, index=True)   # "chats" | "dashboard" | "qchats"
    created_at = Column(DateTime, default=datetime.utcnow)
    
class CounsellorProfile(Base):
    __tablename__ = "counsellor_profiles"
    user_id = Column(Integer, ForeignKey("users.id"), primary_key=True, index=True)
    name = Column(String(255), nullable=False)
    description = Column(Text, nullable=True)
    experience_years = Column(Integer, default=0)
    created_at = Column(DateTime, default=datetime.utcnow)

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String(255), unique=True, index=True, nullable=False)
    password_hash = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    is_counsellor = Column(Integer, default=0)


class Conversation(Base):
    __tablename__ = "conversations"
    id = Column(String(36), primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)
    title = Column(String(255))
    created_at = Column(DateTime, default=datetime.utcnow)

class UserAccess(Base):
    __tablename__ = "user_access"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)          # owner
    counsellor_id = Column(Integer, ForeignKey("users.id"), index=True)    # counsellor
    allow_chats = Column(Integer, default=0)
    allow_q_chats = Column(Integer, default=0)
    allow_dashboard = Column(Integer, default=0)
    created_at = Column(DateTime, default=datetime.utcnow)

class Message(Base):
    __tablename__ = "messages"
    id = Column(Integer, primary_key=True, index=True)
    conversation_id = Column(String(36), ForeignKey("conversations.id"), nullable=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)
    role = Column(String(16))  # 'user' or 'assistant'
    content = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)


class Document(Base):
    __tablename__ = "documents"
    id = Column(String(36), primary_key=True, index=True)
    uploader = Column(String(255))
    filename = Column(String(255))
    uploaded_at = Column(DateTime, default=datetime.utcnow)


class QuestionnaireAttempt(Base):
    __tablename__ = "q_attempts"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)
    day = Column(String(10), index=True)  # yyyy-mm-dd
    score = Column(Float)
    comment = Column(Text)
    tips_json = Column(Text)  # JSON list of strings
    session_id = Column(String(36))  # link to its conversation (optional)
    created_at = Column(DateTime, default=datetime.utcnow)


# NEW: QuestionnaireSession to manage target question count and active flag
class QuestionnaireSession(Base):
    __tablename__ = "q_sessions"
    session_id = Column(String(36), primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), index=True)
    target_questions = Column(Integer, default=MIN_QUESTIONS)
    is_active = Column(Integer, default=1)  # 1 = active, 0 = closed
    created_at = Column(DateTime, default=datetime.utcnow)


Base.metadata.create_all(bind=engine)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ---------------- Auth utils ----------------
pwd_ctx = CryptContext(schemes=["bcrypt"], deprecated="auto")
bearer = HTTPBearer()


def hash_password(pw: str) -> str:
    return pwd_ctx.hash(pw)


def verify_password(pw: str, hashed: str) -> bool:
    return pwd_ctx.verify(pw, hashed)


def create_access_token(sub: str, minutes: int = JWT_EXPIRE_MINUTES) -> str:
    exp = datetime.utcnow() + timedelta(minutes=minutes)
    return jwt.encode({"sub": sub, "exp": exp}, SECRET_KEY, algorithm="HS256")


def decode_token(token: str) -> str:
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        return payload.get("sub")
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")


# Admin principal helper
class AdminPrincipal:
    id = -1
    email = "admin"
    is_admin = True


def is_admin_principal(obj: Any) -> bool:
    return getattr(obj, "email", None) == "admin" or getattr(obj, "is_admin", False) is True


def get_current_user(
    creds: HTTPAuthorizationCredentials = Depends(bearer), db: Session = Depends(get_db)
):
    token = creds.credentials
    sub = decode_token(token)
    # Hardcoded admin principal
    if sub == "admin":
        return AdminPrincipal()
    # Normal user
    user = db.query(User).filter(User.email == sub).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user

def _html_to_text(html: str) -> str:
    # Very small fallback: strip tags
    return re.sub(r"<[^>]+>", "", html or "").strip()

def send_email(to_email: str, subject: str, html_body: str, text_body: Optional[str] = None) -> bool:
    """Send email with optional TLS/SSL and auth. Returns True on success (raises on hard errors)."""
    msg = MIMEMultipart("alternative")
    # normalize "From" (extract address, keep display name)
    from_name, from_addr = parseaddr(SMTP_FROM)
    if not from_addr:
        from_addr = "no-reply@psychsupport.local"
    msg["From"] = formataddr((from_name or "Psych Support", from_addr))
    msg["To"] = to_email
    msg["Subject"] = subject

    part_text = MIMEText(text_body or _html_to_text(html_body), "plain", "utf-8")
    part_html = MIMEText(html_body, "html", "utf-8")
    msg.attach(part_text)
    msg.attach(part_html)

    context = ssl.create_default_context()
    try:
        if SMTP_SECURE == "ssl":
            server = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=SMTP_TIMEOUT, context=context)
        else:
            server = smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=SMTP_TIMEOUT)

        server.ehlo()
        if SMTP_SECURE == "starttls":
            server.starttls(context=context)
            server.ehlo()

        if SMTP_USER:
            server.login(SMTP_USER, SMTP_PASSWORD or "")

        server.sendmail(from_addr, [to_email], msg.as_string())
        server.quit()
        return True
    except Exception as e:
        # Surface clear info to logs; callers already catch per-recipient in your code.
        print(f"[MAIL] send_email to={to_email} failed:", repr(e))
        raise

        
def users_without_any_attempts(db: Session) -> list[User]:
    # users having 0 rows in q_attempts
    return (
        db.query(User)
        .outerjoin(QuestionnaireAttempt, QuestionnaireAttempt.user_id == User.id)
        .group_by(User.id)
        .having(func.count(QuestionnaireAttempt.id) == 0)
        .all()
    )

def users_with_attempts_on_day(db: Session, day_str: str) -> list[User]:
    sub = db.query(QuestionnaireAttempt.user_id).filter(QuestionnaireAttempt.day == day_str).distinct().subquery()
    return db.query(User).filter(User.id.in_(sub)).all()

def attempts_for_user_on_day(db: Session, user_id: int, day_str: str) -> list[QuestionnaireAttempt]:
    return (
        db.query(QuestionnaireAttempt)
        .filter(QuestionnaireAttempt.user_id == user_id, QuestionnaireAttempt.day == day_str)
        .order_by(QuestionnaireAttempt.created_at.asc())
        .all()
    )

def attempts_for_user_all(db: Session, user_id: int) -> list[QuestionnaireAttempt]:
    return (
        db.query(QuestionnaireAttempt)
        .filter(QuestionnaireAttempt.user_id == user_id)
        .order_by(QuestionnaireAttempt.created_at.asc())
        .all()
    )

def compose_reminder_html(username: str) -> tuple[str, str]:
    subject = "Gentle reminder: Try today’s quick wellness check-in"
    html = f"""
    <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
      <h2>Hi {username},</h2>
      <p>This is a quick nudge to try our <b>2–3 minute</b> mental wellness questionnaire.</p>
      <ul>
        <li>Track your mood & stress</li>
        <li>Get personalized tips</li>
      </ul>
      <p>You can start it from your dashboard any time.</p>
      <p style="color:#666">If you didn’t sign up for this, you can ignore this mail.</p>
    </div>
    """
    return subject, html

def compose_daily_summary_html(username: str, attempts: list[QuestionnaireAttempt]) -> tuple[str, str]:
    scores = [a.score for a in attempts if a.score is not None]
    avg = sum(scores) / len(scores) if scores else 0.0
    latest_comment = attempts[-1].comment if attempts and attempts[-1].comment else ""
    subject = "Your daily check-in summary"
    html = f"""
    <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
      <h2>Hi {username}, here’s your today’s summary</h2>
      <p><b>Attempts today:</b> {len(attempts)}<br/>
         <b>Average score:</b> {avg:.1f}/100</p>
      {"<p><b>Latest note:</b> " + latest_comment + "</p>" if latest_comment else ""}
      <p>Keep going—small, steady steps make a difference.</p>
    </div>
    """
    return subject, html

def compose_periodic_html(username: str, attempts: list[QuestionnaireAttempt]) -> tuple[str, str]:
    if not attempts:
        return ("Your check-in summary", f"<div style='font-family:system-ui,Segoe UI,Arial,sans-serif'><p>No attempts yet.</p></div>")

    scores = [a.score for a in attempts if a.score is not None]
    avg_all = sum(scores)/len(scores) if scores else 0.0
    max_a = max(attempts, key=lambda a: (a.score or 0))
    min_a = min(attempts, key=lambda a: (a.score or 0))

    # Monthly groups
    monthly: dict[str, list[QuestionnaireAttempt]] = {}
    for a in attempts:
        ym = (a.day or a.created_at.date().isoformat())[:7]  # YYYY-MM
        monthly.setdefault(ym, []).append(a)

    monthly_html = []
    for ym, rows in sorted(monthly.items(), reverse=True):
        s = [r.score for r in rows if r.score is not None]
        avg_m = sum(s)/len(s) if s else 0.0
        # list scores compactly
        scores_str = ", ".join(f"{(r.score or 0):.0f}" for r in rows)
        monthly_html.append(f"""
          <div style="margin-bottom:10px">
            <h4 style="margin:0">Month {ym}</h4>
            <div>Attempts: {len(rows)} &middot; Avg: {avg_m:.1f}</div>
            <div>Scores: {scores_str}</div>
          </div>
        """)
    subject = "Your check-in progress summary"
    html = f"""
    <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
      <h2>Hi {username}, here’s your progress</h2>
      <p><b>Total attempts:</b> {len(attempts)}<br/>
         <b>Average score:</b> {avg_all:.1f}/100</p>
      <p><b>Highest:</b> {max_a.score:.1f} — {max_a.comment or "—"}</p>
      <p><b>Lowest:</b> {min_a.score:.1f} — {min_a.comment or "—"}</p>
      <hr/>
      <h3 style="margin-top:10px">By month</h3>
      {''.join(monthly_html)}
    </div>
    """
    return subject, html

def send_reminders_to_incomplete_users(db: Session) -> int:
    users = users_without_any_attempts(db)
    sent = 0
    for u in users:
        try:
            subject, html = compose_reminder_html(username_from_email(u.email))
            send_email(u.email, subject, html)
            sent += 1
        except Exception as e:
            print("[MAIL] reminder failed for", u.email, e)
    return sent

def send_daily_summaries_today(db: Session) -> int:
    today = date.today().isoformat()
    users = users_with_attempts_on_day(db, today)
    sent = 0
    for u in users:
        rows = attempts_for_user_on_day(db, u.id, today)
        if not rows:
            continue
        try:
            subject, html = compose_daily_summary_html(username_from_email(u.email), rows)
            send_email(u.email, subject, html)
            sent += 1
        except Exception as e:
            print("[MAIL] daily summary failed for", u.email, e)
    return sent

def send_periodic_summaries_all(db: Session) -> int:
    users = db.query(User).all()
    sent = 0
    for u in users:
        rows = attempts_for_user_all(db, u.id)
        if not rows:
            continue
        try:
            subject, html = compose_periodic_html(username_from_email(u.email), rows)
            send_email(u.email, subject, html)
            sent += 1
        except Exception as e:
            print("[MAIL] periodic summary failed for", u.email, e)
    return sent

def require_admin(user=Depends(get_current_user)):
    if is_admin_principal(user):
        return user
    raise HTTPException(status_code=403, detail="Admin only")


# ---------------- RAG ----------------
_qdrant: Optional[QdrantClient] = None
_embedder: Optional[SentenceTransformer] = None

def init_embedder() -> Optional[SentenceTransformer]:
    global _embedder
    if _embedder is not None:
        return _embedder
    try:
        _embedder = SentenceTransformer(EMBED_MODEL)
        print("[RAG] Embedder loaded:", EMBED_MODEL)
        return _embedder
    except Exception as e:
        print("[RAG] Embedder load failed:", e)
        _embedder = None
        return None


def init_qdrant() -> Optional[QdrantClient]:
    global _qdrant
    if _qdrant is not None:
        return _qdrant
    try:
        _qdrant = QdrantClient(url=QDRANT_URL)
        try:
            _qdrant.get_collection(COLLECTION)
        except Exception:
            emb = init_embedder()
            if emb is None:
                raise RuntimeError("Cannot create collection because embedder failed")
            vec_size = emb.get_sentence_embedding_dimension()
            _qdrant.create_collection(
                collection_name=COLLECTION,
                vectors_config=qmodels.VectorParams(size=vec_size, distance=qmodels.Distance.COSINE),
            )
        # ✅ Optional but recommended: index user_id for fast filtering
        try:
            _qdrant.create_payload_index(
                collection_name=COLLECTION,
                field_name="user_id",
                field_schema=qmodels.PayloadSchemaType.INTEGER,
            )
        except Exception:
            # index may already exist; ignore
            pass

        print("[RAG] Qdrant connected:", QDRANT_URL)
        return _qdrant
    except Exception as e:
        print("[RAG] Qdrant init failed:", e)
        _qdrant = None
        return None


def embed_texts(texts: List[str]) -> List[List[float]]:
    emb = init_embedder()
    if emb is None:
        raise RuntimeError("Embedder not available")
    vectors = emb.encode(texts, normalize_embeddings=True)
    return [v.tolist() if hasattr(v, "tolist") else list(v) for v in vectors]


def upsert_documents_to_qdrant(payloads: List[dict]) -> int:
    qc = init_qdrant()
    if qc is None:
        raise RuntimeError("Qdrant not available")
    vectors = embed_texts([p["text"] for p in payloads])
    points = []
    for i, p in enumerate(payloads):
        points.append(
            qmodels.PointStruct(
                id=p["id"],
                vector=vectors[i],
                payload={
                    "text": p["text"],
                    "meta": p.get("meta", {}),
                    "user_id": p.get("user_id")  # ✅ add user_id here
                },
            )
        )
    qc.upsert(collection_name=COLLECTION, points=points)
    return len(points)


def search_qdrant(query: str, limit: int = TOP_K, user_id: Optional[int] = None) -> List[dict]:
    qc = init_qdrant()
    if qc is None:
        return []
    qvec = embed_texts([query])[0]


    # Depending on qdrant-client version, the kwarg is either `query_filter` or `filter`.
    # Most recent clients use `query_filter=`.
    hits = qc.search(
        collection_name=COLLECTION,
        query_vector=qvec,
        limit=limit  # <-- if this errors, change to `filter=q_filter`
    )
    results = []
    for h in hits:
        results.append({
            "id": str(h.id),
            "text": h.payload.get("text", ""),
            "meta": h.payload.get("meta", {}),
            "score": float(h.score) if h.score is not None else 0.0,
        })
    return results


# ---------------- helpers: sources as excerpts ----------------
SOURCES_MARKER_PREFIX = "<!--SOURCES_JSON:"
SOURCES_MARKER_SUFFIX = "-->"

def build_source_items(contexts: List[dict]) -> List[dict]:
    """(Legacy helper) Not used for rendering anymore; kept for compatibility."""
    items = []
    for c in contexts:
        meta = c.get("meta") or {}
        fname = meta.get("filename") or meta.get("uploader") or c.get("id")
        text = (c.get("text") or "").strip().replace("\n", " ")
        excerpt = text[:280] + ("…" if len(text) > 280 else "")
        items.append({"filename": fname, "excerpt": excerpt})
    return items

# -------------- NEW helpers for numbered citations & OOC prefix --------------

_YT_PATTERNS = [
    r"(?:https?://)?(?:www\.)?youtube\.com/watch\?v=([A-Za-z0-9_\-]{6,})",
    r"(?:https?://)?youtu\.be/([A-Za-z0-9_\-]{6,})",
    r"(?:https?://)?(?:www\.)?youtube\.com/embed/([A-Za-z0-9_\-]{6,})",
]

def normalize_youtube_url(url: str) -> str:
    url = (url or "").strip()
    for pat in _YT_PATTERNS:
        m = re.match(pat, url)
        if m:
            vid = m.group(1)
            return f"https://www.youtube.com/watch?v={vid}"
    # fallback: keep as-is if it's http(s) and looks like youtube domain
    if "youtube.com" in url or "youtu.be" in url:
        return url
    raise HTTPException(status_code=400, detail="Only YouTube links are supported for now.")

def _tags_to_str(tags: Optional[List[str]]) -> str:
    if not tags: return ""
    return ",".join(t.strip() for t in tags if t and t.strip())

def _str_to_tags(s: Optional[str]) -> List[str]:
    if not s: return []
    return [t.strip() for t in s.split(",") if t.strip()]


OOC_PREFIX = "🌐 I couldn’t find context in the uploaded sources — this answer is drawn from general knowledge!!"


def is_out_of_context(contexts: List[dict]) -> bool:
    best_score = max((c.get("score", 0.0) for c in contexts), default=0.0)
    return (best_score < OUT_OF_CONTEXT_THRESHOLD) or (not contexts)

def _number_contexts(contexts: List[dict]) -> List[dict]:
    """Return contexts[:TOP_K] with an added 1-based 'index' key."""
    ordered = []
    for i, c in enumerate(contexts[:TOP_K], start=0):
        cc = dict(c)
        cc["index"] = i
        ordered.append(cc)
    return ordered

def build_prompt_with_numbered_context(user_msg: str, contexts: List[dict]) -> tuple[str, List[dict], bool]:
    """
    Returns: (prompt, numbered_contexts, out_of_context_flag)
    - numbered_contexts: contexts[:TOP_K] with .index
    - prompt instructs the model to cite using [n] properly and to emit OOC_PREFIX if out-of-context
    """
    system_prompt = (
        "You are a compassionate, knowledgeable mental wellness assistant. "
        "Prefer concise, practical guidance. "
        "When you use information from a provided source, append the citation marker like [1], [2] "
        "exactly matching the numbered CONTEXT SOURCES below. "
        "Do not invent citation numbers. If a statement is general knowledge, do not cite it."
    )

    numbered = _number_contexts(contexts)
    ooc = is_out_of_context(contexts)

    if ooc:
        # Explicitly instruct the first line with emoji
        prompt = (
            f"SYSTEM:\n{system_prompt}\n\n"
            f"USER:\n{user_msg}\n\n"
            "ASSISTANT:\n"
            f"Begin your response with EXACTLY this line followed by a blank line:\n"
            f"\"{OOC_PREFIX}\"\n\n"
            "Then provide a clear, supportive answer based on general knowledge only. "
            "Do not add any [n] citations because no uploaded context is being used."
        )
        return prompt, numbered, True

    # Build a compact numbered context block
    blocks = []
    for c in numbered:
        meta = c.get("meta") or {}
        src = meta.get("filename") or meta.get("uploader") or c.get("id")
        text = (c.get("text") or "")[:800]
        score = float(c.get("score") or 0.0)
        blocks.append(f"[{c['index']}] filename={src} score={score:.4f}\n{text}")
    ctx = "\n\n".join(blocks)

    prompt = (
        f"SYSTEM:\n{system_prompt}\n\n"
        f"CONTEXT SOURCES (numbered):\n{ctx}\n\n"
        f"USER:\n{user_msg}\n\n"
        "ASSISTANT:\nUse the numbered sources above where appropriate; append [n] after the relevant sentence. "
        "Only cite numbers that exist in the context. If a statement is not supported by the context, don't cite."
    )
    return prompt, numbered, False


def append_cited_sources_marker(reply_text: str, numbered_contexts: List[dict]) -> str:
    """
    Scan reply for [n] citations and attach ONLY those sources as JSON in the marker.
    Keeps order of first appearance. Each item: {index, excerpt, filename, id}
    """
    if not reply_text or not numbered_contexts:
        return reply_text

    pat_num = re.compile(r"\[(\d+)\]")
    pat_legacy = re.compile(r"\[source(\d+)\]", re.IGNORECASE)

    seen: List[int] = []
    for m in pat_num.finditer(reply_text):
        i = int(m.group(1))
        if i not in seen:
            seen.append(i)
    for m in pat_legacy.finditer(reply_text):
        i = int(m.group(1))
        if i not in seen:
            seen.append(i)

    if not seen:
        return reply_text

    by_index = {c["index"]: c for c in numbered_contexts}
    items = []
    for idx in seen:
        c = by_index.get(idx)
        if not c:
            continue
        text = (c.get("text") or "").strip().replace("\n", " ")
        excerpt = text[:280] + ("…" if len(text) > 280 else "")
        meta = c.get("meta") or {}
        fname = meta.get("filename") or meta.get("uploader") or c.get("id")
        items.append({
            "index": idx,
            "excerpt": excerpt,
            "filename": fname,
            "id": c.get("id")
        })

    if not items:
        return reply_text

    payload = json.dumps(items, ensure_ascii=False)
    return f"{reply_text}\n\n{SOURCES_MARKER_PREFIX}{payload}{SOURCES_MARKER_SUFFIX}"


def split_reply_and_sources_for_api(text: str) -> tuple[str, List[dict]]:
    """Split assistant reply and extract JSON source items from marker."""
    if not text:
        return text, []
    start = text.rfind(SOURCES_MARKER_PREFIX)
    if start == -1:
        return text, []
    end = text.find(SOURCES_MARKER_SUFFIX, start)
    if end == -1:
        return text, []
    try:
        payload = text[start + len(SOURCES_MARKER_PREFIX): end]
        items = json.loads(payload)
        clean = text[:start].rstrip()
        return clean, items if isinstance(items, list) else []
    except Exception:
        return text, []


# ---------------- LLMs ----------------
def call_ollama(prompt: str, num_predict: int = 512) -> str:
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/generate",
            json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False, "options": {"num_predict": num_predict}},
            timeout=600,
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("response") or data.get("text") or ""
    except requests.exceptions.Timeout:
        return "The local model (Mistral) timed out. For big questions, switch provider to Groq."
    except Exception as e:
        print("[LLM] Ollama call failed:", e)
        return "Local model is not available right now."

def extract_text_from_file(file: UploadFile) -> str:
    name = file.filename.lower()
    if name.endswith(".txt"):
        return file.file.read().decode("utf-8", errors="ignore")
    elif name.endswith(".pdf"):
        from PyPDF2 import PdfReader
        pdf = PdfReader(file.file)
        return "\n".join([p.extract_text() or "" for p in pdf.pages])
    elif name.endswith(".docx"):
        import docx
        doc = docx.Document(file.file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif name.endswith((".xls", ".xlsx")):
        df = pd.read_excel(file.file)
        return "\n".join(df.astype(str).fillna("").apply(lambda r: " ".join(r), axis=1))
    elif name.endswith(".csv"):
        df = pd.read_csv(file.file)
        return "\n".join(df.astype(str).fillna("").apply(lambda r: " ".join(r), axis=1))
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

def call_groq(prompt: str, num_predict: int = 512) -> str:
    if not _groq_client:
        return "Groq API key not configured. Add GROQ_API_KEY to .env."
    try:
        chat_completion = _groq_client.chat.completions.create(
            model=GROQ_MODEL_DEFAULT,
            messages=[
                {"role": "system", "content": "You are a helpful, compassionate assistant."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=num_predict,
            temperature=0.2,
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        print("[LLM] Groq SDK call failed:", e)
        return "Groq API is not available right now."
    


# ---------------- FastAPI ----------------
app = FastAPI(title="Psych Support - RAG Chat")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOWED_ORIGINS,   # ["*"]
    allow_credentials=False,              # <- change this if using "*"
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------- Schemas ----------------
class CounsellorAppointmentCreate(BaseModel):
    user_id: int
    start_time: datetime
    duration_minutes: int
    meeting_link: str

class AppointmentClose(BaseModel):
    report: str
    
class AppointmentDecision(BaseModel):
    status: Literal["accepted", "rejected"]
    meeting_link: Optional[str] = None
    
class AppointmentCreate(BaseModel):
    counsellor_id: int
    start_time: datetime
    duration_minutes: int
    
class AccessRequestIn(BaseModel):
    user_id: int
    request_type: Literal["chats", "dashboard"]

class VideoCreate(BaseModel):
    title: str
    url: str
    description: Optional[str] = None
    tags: List[str] = []
    is_public: bool = True
    target_user_id: Optional[int] = None
    
class CounsellorUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    experience_years: Optional[int] = None

    
class CounsellorCreate(BaseModel):
    user_id: int
    name: str
    description: Optional[str] = None
    experience_years: int = 0

class VideoOut(BaseModel):
    id: str
    title: str
    url: str
    platform: str
    description: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    is_public: bool
    # NEW:
    added_by: Optional[str] = None
    access: str
    created_at: datetime



class RegisterIn(BaseModel):
    email: EmailStr
    password: str

# CHANGED: allow username OR email for login
class LoginIn(BaseModel):
    identifier: str  # email OR "admin"
    password: str
    
class AccessCreate(BaseModel):
    counsellor_id: int
    allow_chats: bool = False
    allow_q_chats: bool = False
    allow_dashboard: bool = False

class AccessUpdate(BaseModel):
    allow_chats: Optional[bool] = None
    allow_q_chats: Optional[bool] = None
    allow_dashboard: Optional[bool] = None

class ChatIn(BaseModel):
    message: str
    session_id: Optional[str] = None
    num_predict: Optional[int] = 512
    provider: Literal["mistral", "groq"] = "mistral"

class ChatOut(BaseModel):
    reply: str
    sources: List[dict] = Field(default_factory=list)

class UpdateTitleIn(BaseModel):
    title: str

# Questionnaire
class QStartOut(BaseModel):
    session_id: str
    question: str

class QAnswerIn(BaseModel):
    session_id: str
    answer: str

class QAnswerOut(BaseModel):
    done: bool
    question: Optional[str] = None
    score: Optional[float] = None
    comment: Optional[str] = None
    tips: List[str] = Field(default_factory=list)



def delete_user_vectors(user_id: int) -> None:
    qc = init_qdrant()
    if not qc:
        return
    try:
        qc.delete(
            collection_name=COLLECTION,
            points_selector=qmodels.FilterSelector(
                filter=qmodels.Filter(
                    must=[qmodels.FieldCondition(key="user_id", match=qmodels.MatchValue(value=user_id))]
                )
            ),
            wait=True,
        )
        print(f"[Qdrant] Deleted vectors for user_id={user_id}")
    except Exception as e:
        print("[Qdrant] delete_user_vectors failed:", e)



def short_title(text: str) -> str:
    t = (text or "").strip().replace("\n", " ")
    if not t: return "New Chat"
    words = t.split()
    s = " ".join(words[:8])
    if len(words) > 8: s += "…"
    return s[:80]

def username_from_email(email: str) -> str:
    return (email or "").split("@", 1)[0] or "there"



# ---------------- Routes: meta ----------------
@app.get("/health")
def health(user=Depends(get_current_user)):
    return {
        "ok": True,
        "qdrant": init_qdrant() is not None,
        "embedder": init_embedder() is not None,
        "me": {
            "email": user.email if not is_admin_principal(user) else "admin",
            "username": "admin" if is_admin_principal(user) else username_from_email(user.email),
            "is_admin": is_admin_principal(user),
            "is_counsellor": bool(getattr(user, "is_counsellor", 0)),
        },
        "provider_defaults": {"mistral": OLLAMA_MODEL, "groq": GROQ_MODEL_DEFAULT if GROQ_API_KEY else None},
    }


# ---------------- Auth ----------------
@app.post("/auth/register")
def register(payload: RegisterIn, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.email == payload.email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered. Please log in instead.")
    user = User(email=payload.email, password_hash=hash_password(payload.password))
    db.add(user); db.commit(); db.refresh(user)
    token = create_access_token(user.email)
    return {"access_token": token, "token_type": "bearer", "is_admin": False, "is_counsellor": bool(user.is_counsellor),"user_id": user.id}

@app.post("/auth/login")
def login(payload: LoginIn, db: Session = Depends(get_db)):
    # Hardcoded admin login
    if payload.identifier == "admin" and payload.password == "admin":
        token = create_access_token("admin")
        return {"access_token": token, "token_type": "bearer", "is_admin": True, "is_counsellor": False, "user_id": None}

    # Normal user flow
    user = db.query(User).filter(User.email == payload.identifier).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email/username or password")
    token = create_access_token(user.email)
    return {
        "access_token": token,
        "token_type": "bearer",
        "is_admin": False,
        "is_counsellor": bool(getattr(user, "is_counsellor", False)),
        "user_id": user.id,
    }

    # Normal user flow
    user = db.query(User).filter(User.email == payload.identifier).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email/username or password")
    token = create_access_token(user.email)
    return {"access_token": token, "token_type": "bearer", "is_admin": False}

# Aliases
@app.post("/register")
def register_alias(payload: RegisterIn, db: Session = Depends(get_db)):
    return register(payload, db)

@app.post("/login")
def login_alias(payload: LoginIn, db: Session = Depends(get_db)):
    return login(payload, db)


# ---------------- Sessions/Chats ----------------
@app.get("/sessions")
def list_sessions(user=Depends(get_current_user), db: Session = Depends(get_db)):
    # Admin has no personal sessions; return empty list
    if is_admin_principal(user):
        return []
    rows = db.query(Conversation).filter(Conversation.user_id == user.id).order_by(Conversation.created_at.desc()).all()
    return [{"id": r.id, "title": r.title, "created_at": r.created_at.isoformat()} for r in rows]

@app.get("/sessions/{session_id}/messages")
def session_messages(session_id: str, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admins don't have chat history.")
    msgs = (
        db.query(Message)
        .filter(Message.conversation_id == session_id, Message.user_id == user.id)
        .order_by(Message.created_at.asc())
        .all()
    )
    return [{"role": m.role, "content": m.content, "created_at": m.created_at.isoformat()} for m in msgs]

@app.patch("/sessions/{session_id}")
def update_session_title(session_id: str, payload: UpdateTitleIn, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admins cannot rename user sessions.")
    conv = db.query(Conversation).filter(Conversation.id == session_id, Conversation.user_id == user.id).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Session not found")
    conv.title = payload.title[:255]
    db.commit()
    return {"id": conv.id, "title": conv.title}


# ---------------- Ingest ----------------
@app.post("/ingest")
def ingest(file: UploadFile = File(...), user=Depends(get_current_user), db: Session = Depends(get_db)):
    raw = extract_text_from_file(file)
    chunks = [raw[i:i+800].strip() for i in range(0, len(raw), 800) if raw[i:i+800].strip()]
    payloads = [{
        "id": str(uuid.uuid4()),
        "text": c,
        "meta": {"uploader": user.email, "filename": file.filename},
        "user_id": user.id,
    } for c in chunks]
    db.add(Document(id=str(uuid.uuid4()), uploader=user.email, filename=file.filename))
    db.commit()
    count = upsert_documents_to_qdrant(payloads)
    return {"ok": True, "ingested_chunks": count}


# ---------------- Chat (updated: numbered citations + emoji OOC prefix) ----------------
@app.post("/chat", response_model=ChatOut)
def chat(body: ChatIn, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admin cannot chat.")
    # RAG search
    try:
        contexts = search_qdrant(body.message, limit=TOP_K, user_id=user.id)
    except Exception as e:
        print("[RAG] search error:", e, traceback.format_exc())
        contexts = []

    prompt, numbered_ctx, ooc = build_prompt_with_numbered_context(body.message, contexts)
    max_toks = body.num_predict or 512
    raw_reply = call_groq(prompt, num_predict=max_toks) if body.provider == "groq" else call_ollama(prompt, num_predict=max_toks)

    # Guarantee first line for out-of-context
    reply = (raw_reply or "").strip()
    if ooc and not reply.startswith(OOC_PREFIX):
        reply = f"\n\n{reply}"

    # Attach ONLY cited sources (based on [n] in reply)
    reply_with_marker = append_cited_sources_marker(reply, numbered_ctx)

    # Create a conversation ONLY NOW (first message), if missing
    conv_id = body.session_id
    if not conv_id:
        conv = Conversation(id=str(uuid.uuid4()), user_id=user.id, title=short_title(body.message))
        db.add(conv); db.commit(); conv_id = conv.id

    # save
    try:
        db.add_all([
            Message(conversation_id=conv_id, user_id=user.id, role="user", content=body.message),
            Message(conversation_id=conv_id, user_id=user.id, role="assistant", content=reply_with_marker),
        ])
        db.commit()
    except Exception as e:
        print("[DB] save failed:", e, traceback.format_exc())

    # Build sources payload for response model (only cited)
    _, items = split_reply_and_sources_for_api(reply_with_marker)
    return {"reply": reply_with_marker, "sources": items}


@app.post("/chat_stream")
def chat_stream(body: ChatIn, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admin cannot chat.")
    def gen():
        # RAG search
        try:
            contexts = search_qdrant(body.message, limit=TOP_K, user_id=user.id)
        except Exception as e:
            print("[RAG] search error:", e, traceback.format_exc())
            contexts = []

        prompt, numbered_ctx, ooc = build_prompt_with_numbered_context(body.message, contexts)
        max_toks = body.num_predict or 512

        chunks: List[str] = []

        try:
            if body.provider == "groq":
                if not _groq_client:
                    yield json.dumps({"type":"delta","text":"[Groq key missing]\n"}) + "\n"
                else:
                    stream = _groq_client.chat.completions.create(
                        model=GROQ_MODEL_DEFAULT,
                        messages=[
                            {"role":"system","content":"You are a helpful, compassionate assistant."},
                            {"role":"user","content":prompt}
                        ],
                        max_tokens=max_toks, temperature=0.2, stream=True,
                    )
                    for chunk in stream:
                        d = ""
                        if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                            d = chunk.choices[0].delta.content
                        if d:
                            chunks.append(d)
                            yield json.dumps({"type":"delta","text": d}) + "\n"
            else:
                with requests.post(
                    f"{OLLAMA_URL}/api/generate",
                    json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": True, "options": {"num_predict": max_toks}},
                    stream=True, timeout=600,
                ) as r:
                    r.raise_for_status()
                    for line in r.iter_lines(decode_unicode=True):
                        if not line: continue
                        try:
                            ev = json.loads(line)
                        except Exception:
                            continue
                        d = ev.get("response")
                        if d:
                            chunks.append(d)
                            yield json.dumps({"type":"delta","text": d}) + "\n"
        except requests.exceptions.Timeout:
            yield json.dumps({"type":"delta","text":"[Provider timed out]\n"}) + "\n"
        finally:
            body_text = "".join(chunks).strip()
            final_text = body_text
            # If model ignored instruction, ensure the first line is present
            if ooc and not final_text.startswith(OOC_PREFIX):
                final_text = f"\n\n{final_text}"

            final_text_with_marker = append_cited_sources_marker(final_text, numbered_ctx)

            # Create conversation only now (first message), if missing
            conv_id = body.session_id
            if not conv_id:
                conv = Conversation(id=str(uuid.uuid4()), user_id=user.id, title=short_title(body.message))
                db.add(conv); db.commit(); conv_id = conv.id

            try:
                db.add_all([
                    Message(conversation_id=conv_id, user_id=user.id, role="user", content=body.message),
                    Message(conversation_id=conv_id, user_id=user.id, role="assistant", content=final_text_with_marker),
                ])
                db.commit()
            except Exception as e:
                print("[DB] save failed (stream):", e, traceback.format_exc())

            # Send any trailing text (e.g., sources marker that wasn't streamed)
            yield json.dumps({"type":"delta","text": final_text_with_marker[len(final_text):]}) + "\n"
            yield json.dumps({"type":"done"}) + "\n"

    return StreamingResponse(gen(), media_type="application/jsonl")

# ---------------- Chat Delete By USer----------------
@app.get("/videos", response_model=List[VideoOut])
def list_videos(q: Optional[str] = None, db: Session = Depends(get_db), user=Depends(get_current_user)):
    qry = db.query(Video)

    if not is_admin_principal(user):
        if getattr(user, "is_counsellor", 0):
            qry = qry.filter(
                or_(
                    and_(Video.access == "all", Video.is_public == 1),
                    Video.added_by == user.email,
                    Video.access.like("user:%")
                )
            )
        else:
            qry = qry.filter(
                or_(
                    and_(Video.access == "all", Video.is_public == 1),
                    Video.access == f"user:{user.id}"
                )
            )

    if q:
        like = f"%{q.strip()}%"
        qry = qry.filter(or_(
            Video.title.ilike(like),
            Video.description.ilike(like),
            Video.tags.ilike(like)
        ))

    rows = qry.order_by(desc(Video.created_at)).all()
    return [{
        "id": v.id,
        "title": v.title,
        "url": v.url,
        "platform": v.platform,
        "description": v.description,
        "tags": _str_to_tags(v.tags),
        "is_public": bool(v.is_public),
        "added_by": v.added_by,
        "access": v.access,
        "created_at": v.created_at,
    } for v in rows]


@app.post("/admin/emails/reminders")
def admin_send_reminders(db: Session = Depends(get_db), admin=Depends(require_admin)):
    count = send_reminders_to_incomplete_users(db)
    return {"ok": True, "sent": count}

@app.post("/admin/emails/daily-summaries")
def admin_send_daily_summaries(db: Session = Depends(get_db), admin=Depends(require_admin)):
    count = send_daily_summaries_today(db)
    return {"ok": True, "sent": count}

@app.post("/admin/emails/periodic-summaries")
def admin_send_periodic_summaries(db: Session = Depends(get_db), admin=Depends(require_admin)):
    count = send_periodic_summaries_all(db)
    return {"ok": True, "sent": count}

@app.post("/admin/videos", response_model=VideoOut)
def admin_add_video(body: VideoCreate, db: Session = Depends(get_db), admin=Depends(require_admin)):
    url = normalize_youtube_url(str(body.url))
    v = Video(
        id=str(uuid.uuid4()),
        title=body.title.strip(),
        url=url,
        platform="youtube",
        description=(body.description or "").strip() or None,
        tags=_tags_to_str(body.tags),
        is_public=1 if body.is_public else 0,
        # NEW defaults:
        added_by="admin",
        access="all",
    )
    db.add(v); db.commit(); db.refresh(v)
    return {
        "id": v.id,
        "title": v.title,
        "url": v.url,
        "platform": v.platform,
        "description": v.description,
        "tags": _str_to_tags(v.tags),
        "is_public": bool(v.is_public),
        "added_by": v.added_by,
        "access": v.access,
        "created_at": v.created_at,
    }

@app.on_event("startup")
async def start_email_scheduler():
    if not ENABLE_EMAIL_SCHEDULER or AsyncIOScheduler is None or CronTrigger is None:
        print("[SCHED] Scheduler disabled or APScheduler not installed.")
        return
    try:
        scheduler = AsyncIOScheduler(timezone=LOCAL_TZ)

        def _job_wrapper():
            # Run reminders at trigger time
            db = SessionLocal()
            try:
                n = send_reminders_to_incomplete_users(db)
                print(f"[SCHED] Reminders sent: {n}")
            finally:
                db.close()

        # 2PM and 8PM IST every day
        scheduler.add_job(_job_wrapper, CronTrigger(hour=14, minute=0, timezone=LOCAL_TZ))
        scheduler.add_job(_job_wrapper, CronTrigger(hour=20, minute=0, timezone=LOCAL_TZ))
        scheduler.add_job(lambda: send_appointment_reminders(SessionLocal()), CronTrigger(minute="0", hour="*"))

        scheduler.start()
        print("[SCHED] Email scheduler started (IST 14:00 & 20:00).")
    except Exception as e:
        print("[SCHED] Failed to start:", e)

class VideoUpdate(BaseModel):
    title: Optional[str] = None
    url: Optional[AnyHttpUrl] = None
    description: Optional[str] = None
    tags: Optional[List[str]] = None
    is_public: Optional[bool] = None
    
    
def send_appointment_reminders(db: Session):
    now = datetime.utcnow()
    upcoming = db.query(Appointment).filter(
        Appointment.status == "accepted",
        Appointment.start_time > now,
    ).all()

    sent = 0
    for ap in upcoming:
        delta = ap.start_time - now

        # 24h reminder
        if delta <= timedelta(hours=24) and not ap.reminder_24_sent:
            send_email(ap.user.email, "Appointment Reminder (24h)",
                       f"Reminder: You have an appointment with your counsellor at {ap.start_time}.\nMeeting link: {ap.meeting_link}")
            ap.reminder_24_sent = True
            sent += 1

        # 12h reminder
        if delta <= timedelta(hours=12) and not ap.reminder_12_sent:
            send_email(ap.user.email, "Appointment Reminder (12h)",
                       f"Reminder: You have an appointment with your counsellor at {ap.start_time}.\nMeeting link: {ap.meeting_link}")
            ap.reminder_12_sent = True
            sent += 1

    db.commit()
    return sent

@app.patch("/admin/videos/{video_id}", response_model=VideoOut)
def admin_update_video(video_id: str, body: VideoUpdate, db: Session = Depends(get_db), admin=Depends(require_admin)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v: raise HTTPException(status_code=404, detail="Video not found")

    if body.title is not None: v.title = body.title.strip()
    if body.url is not None: v.url = normalize_youtube_url(str(body.url))
    if body.description is not None: v.description = (body.description or "").strip() or None
    if body.tags is not None: v.tags = _tags_to_str(body.tags)
    if body.is_public is not None: v.is_public = 1 if body.is_public else 0

    db.commit(); db.refresh(v)
    return {
        "id": v.id,
        "title": v.title,
        "url": v.url,
        "platform": v.platform,
        "description": v.description,
        "tags": _str_to_tags(v.tags),
        "is_public": bool(v.is_public),
        "added_by": v.added_by,
        "created_at": v.created_at,
    }

@app.delete("/admin/videos/{video_id}")
def admin_delete_video(video_id: str, db: Session = Depends(get_db), admin=Depends(require_admin)):
    v = db.query(Video).filter(Video.id == video_id).first()
    if not v: raise HTTPException(status_code=404, detail="Video not found")
    db.delete(v); db.commit()
    return {"ok": True}


@app.delete("/admin/users/{user_id}/vectors")
def admin_clear_user_vectors(user_id: int, db: Session = Depends(get_db), admin=Depends(require_admin)):
    delete_user_vectors(user_id)
    return {"ok": True}

@app.delete("/sessions/{session_id}")
def delete_session(
    session_id: str = Path(...),
    user=Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admins cannot delete user sessions.")

    conv = (
        db.query(Conversation)
        .filter(Conversation.id == session_id, Conversation.user_id == user.id)
        .first()
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Session not found")

    # delete messages belonging to this user & conversation
    db.query(Message).filter(
        Message.conversation_id == session_id,
        Message.user_id == user.id
    ).delete()

    # delete the conversation row
    db.delete(conv)
    db.commit()
    return {"ok": True}

# ---------------- Questionnaire logic ----------------

def is_continue_prompt_text(text: str) -> bool:
    """Detect our 'continue or score' prompt without using hidden markers."""
    t = (text or "").strip().lower()
    return ("reply with 'continue' or 'score'" in t) or t.startswith("you've answered")

def is_regular_question_text(text: str) -> bool:
    """Count only real questionnaire questions (end with ? and not a continue-prompt)."""
    if not text:
        return False
    t = text.strip()
    if not t.endswith("?"):
        return False
    return not is_continue_prompt_text(t)

def q_system_prompt() -> str:
    return (
        "You are a licensed-therapist-style assistant creating an adaptive mental-health questionnaire. "
        "Ask concise, empathetic questions. Use the user's previous answers to decide the next question. "
        "Do not give scores until asked. Keep each question short and clear."
    )

def q_history(db: Session, user_id: int, session_id: str) -> List[Dict[str, str]]:
    msgs = (
        db.query(Message)
        .filter(Message.conversation_id == session_id, Message.user_id == user_id)
        .order_by(Message.created_at.asc())
        .all()
    )
    return [{"role": m.role, "content": m.content} for m in msgs]

def _extract_qa_transcript(history_pairs: List[Dict[str, str]]) -> List[str]:
    t = []
    for m in history_pairs:
        if m["role"] == "assistant":
            t.append(f"Q: {m['content'].strip()}")
        else:
            t.append(f"A: {m['content']}")
    return t

def next_question_from_llm(history_pairs: List[Dict[str, str]]) -> str:
    transcript = "\n".join(_extract_qa_transcript(history_pairs))
    prompt = (
        q_system_prompt() +
        "\n\nTRANSCRIPT:\n" + transcript +
        "\n\nProduce ONLY the next question (no preface, no JSON, no commentary)."
    )
    text = call_groq(prompt, num_predict=120) if _groq_client else call_ollama(prompt, num_predict=120)
    # take first line as question
    q = (text or "").strip().split("\n")[0].strip()
    if not q.endswith("?"): q += "?"
    return q

def finalize_score_from_history(history_pairs: List[Dict[str, str]]) -> Dict[str, Any]:
    transcript = "\n".join(_extract_qa_transcript(history_pairs))
    prompt = (
        "Analyze the conversation transcript and output JSON ONLY with keys:\n"
        "{\"score\": <0-100>, \"comment\": \"...\", \"tips\": [\"...\", \"...\"]}\n"
        "Be concise, empathetic, and actionable.\n\nTRANSCRIPT:\n" + transcript + "\n\nJSON:"
    )
    text = call_groq(prompt, num_predict=400) if _groq_client else call_ollama(prompt, num_predict=400)
    try:
        start = text.find("{"); end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            data = json.loads(text[start:end+1])
            data["score"] = float(data.get("score", 0.0))
            if not isinstance(data.get("tips"), list):
                data["tips"] = []
            return data
    except Exception:
        pass
    return {"score": 50.0, "comment": "Neutral mood with mixed stress signals.", "tips": ["Take a short walk", "Practice 5-minute breathing"]}


@app.post("/questionnaire/start", response_model=QStartOut)
def questionnaire_start(user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admin cannot start questionnaires.")
    # Create a new conversation
    conv = Conversation(
        id=str(uuid.uuid4()),
        user_id=user.id,
        title=f"[Questionnaire] {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
    )
    db.add(conv); db.commit()

    # Create q_session row with target = MIN_QUESTIONS
    qs = QuestionnaireSession(session_id=conv.id, user_id=user.id, target_questions=MIN_QUESTIONS, is_active=1)
    db.add(qs); db.commit()

    # First question (no markers)
    first_q = "To get started, how would you describe your overall mood today?"
    db.add(Message(conversation_id=conv.id, user_id=user.id, role="assistant", content=first_q))
    db.commit()

    return {"session_id": conv.id, "question": first_q}



@app.post("/questionnaire/answer", response_model=QAnswerOut)
def questionnaire_answer(body: QAnswerIn, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admin cannot answer questionnaires.")

    # Append user answer
    db.add(Message(conversation_id=body.session_id, user_id=user.id, role="user", content=body.answer))
    db.commit()

    # Fetch session settings
    qs = db.query(QuestionnaireSession).filter(
        QuestionnaireSession.session_id == body.session_id,
        QuestionnaireSession.user_id == user.id
    ).first()
    if not qs or not qs.is_active:
        raise HTTPException(status_code=400, detail="Questionnaire session not active.")

    # Load history (assistant + user)
    hist = q_history(db, user.id, body.session_id)

    # Find last assistant message before this user answer
    last_assistant = None
    for m in reversed(hist[:-1]):  # exclude just-added user message
        if m["role"] == "assistant":
            last_assistant = m["content"]
            break

    # Count actual questions asked so far (assistant messages that look like questions, not continue-prompt)
    actual_q_count = sum(
        1 for m in hist if m["role"] == "assistant" and is_regular_question_text(m["content"])
    )

    # If last message was the continue prompt, interpret user's reply
    if last_assistant and is_continue_prompt_text(last_assistant):
        ans = (body.answer or "").strip().lower()

        wants_continue = any(x in ans for x in ["continue", "yes", "y", "go on", "more", "add"])
        wants_score    = any(x in ans for x in ["score", "finish", "end", "done", "result"])

        if wants_continue and not wants_score:
            # bump target by 3 and ask next adaptive question
            qs.target_questions = (qs.target_questions or MIN_QUESTIONS) + CONTINUE_STEP
            db.commit()
            q = next_question_from_llm(hist)
            db.add(Message(conversation_id=body.session_id, user_id=user.id, role="assistant", content=q))
            db.commit()
            return {"done": False, "question": q}

        if wants_score and not wants_continue:
            # finalize now
            data = finalize_score_from_history(hist)
            score = float(data.get("score", 0.0))
            comment = data.get("comment") or ""
            tips = data.get("tips") or []
            summary_text = f"**Summary**\nScore: {score:.1f}/100\n\n{comment}\n\nTips:\n" + "\n".join([f"- {t}" for t in tips])
            db.add(Message(conversation_id=body.session_id, user_id=user.id, role="assistant", content=summary_text))
            # save attempt
            today = date.today().isoformat()
            db.add(QuestionnaireAttempt(
                user_id=user.id, day=today, score=score, comment=comment, tips_json=json.dumps(tips), session_id=body.session_id
            ))
            qs.is_active = 0
            db.commit()
            return {"done": True, "score": score, "comment": comment, "tips": tips}

        # If unclear answer to continue prompt, politely re-ask the same continue question
        prompt_text = (
            f"You've answered {actual_q_count} questions so far. "
            f"Would you like to continue with {CONTINUE_STEP} more questions, or get your score now? "
            "Reply with 'continue' or 'score'."
        )
        db.add(Message(conversation_id=body.session_id, user_id=user.id, role="assistant", content=prompt_text))
        db.commit()
        return {"done": False, "question": prompt_text}

    # If we've reached target, ask the clear 'continue or score' question (no markers)
    if actual_q_count >= (qs.target_questions or MIN_QUESTIONS):
        prompt_text = (
            f"You've answered {actual_q_count} questions so far. "
            f"Would you like to continue with {CONTINUE_STEP} more questions, or get your score now? "
            "Reply with 'continue' or 'score'."
        )
        db.add(Message(conversation_id=body.session_id, user_id=user.id, role="assistant", content=prompt_text))
        db.commit()
        return {"done": False, "question": prompt_text}

    # Otherwise: ask next adaptive question
    q = next_question_from_llm(hist)
    db.add(Message(conversation_id=body.session_id, user_id=user.id, role="assistant", content=q))
    db.commit()
    return {"done": False, "question": q}


# ---------------- Questionnaire dashboard (fast) ----------------
class QDashItem(BaseModel):
    day: str
    attempts: int
    avg_score: float
    latest_comment: Optional[str] = None

@app.get("/questionnaire/dashboard", response_model=List[QDashItem])
def questionnaire_dashboard(user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        return []
    # AVG + COUNT per day
    agg_rows = (
        db.query(
            QuestionnaireAttempt.day.label("day"),
            func.count(QuestionnaireAttempt.id).label("attempts"),
            func.avg(QuestionnaireAttempt.score).label("avg_score"),
        )
        .filter(QuestionnaireAttempt.user_id == user.id)
        .group_by(QuestionnaireAttempt.day)
        .order_by(desc(QuestionnaireAttempt.day))
        .all()
    )

    # Latest comment per day (using max created_at)
    sub_latest = (
        db.query(
            QuestionnaireAttempt.day.label("day"),
            func.max(QuestionnaireAttempt.created_at).label("max_ts"),
        )
        .filter(QuestionnaireAttempt.user_id == user.id)
        .group_by(QuestionnaireAttempt.day)
        .subquery()
    )

    latest_rows = (
        db.query(QuestionnaireAttempt.day, QuestionnaireAttempt.comment)
        .join(sub_latest, (QuestionnaireAttempt.day == sub_latest.c.day) & (QuestionnaireAttempt.created_at == sub_latest.c.max_ts))
        .filter(QuestionnaireAttempt.user_id == user.id)
        .all()
    )
    latest_map = {d: c for d, c in latest_rows}

    out = []
    for r in agg_rows:
        out.append({
            "day": r.day,
            "attempts": int(r.attempts or 0),
            "avg_score": round(float(r.avg_score or 0.0), 2),
            "latest_comment": latest_map.get(r.day, None),
        })
    return out


# ---------------- Admin Endpoints ----------------
@app.get("/admin/users")
def admin_list_users(db: Session = Depends(get_db), admin=Depends(require_admin)):
    users = db.query(User).order_by(User.created_at.desc()).all()
    out = []
    for u in users:
        chat_cnt = db.query(Conversation).filter(Conversation.user_id == u.id).count()
        msg_cnt = db.query(Message).filter(Message.user_id == u.id).count()
        qa_cnt = db.query(QuestionnaireAttempt).filter(QuestionnaireAttempt.user_id == u.id).count()
        active_qs = db.query(QuestionnaireSession).filter(
            QuestionnaireSession.user_id == u.id,
            QuestionnaireSession.is_active == 1
        ).count()

        # 🔹 Determine role
        role = "counsellor" if getattr(u, "is_counsellor", 0) else "user"

        out.append({
            "id": u.id,
            "email": u.email,
            "created_at": u.created_at.isoformat(),
            "conversations": chat_cnt,
            "messages": msg_cnt,
            "questionnaire_attempts": qa_cnt,
            "active_questionnaires": active_qs,
            "role": role
        })
    return out



@app.delete("/admin/users/{user_id}")
def admin_delete_user(user_id: int, db: Session = Depends(get_db), admin=Depends(require_admin)):
    # ✅ delete their vectors first
    delete_user_vectors(user_id)

    db.query(Message).filter(Message.user_id == user_id).delete()
    db.query(Conversation).filter(Conversation.user_id == user_id).delete()
    db.query(QuestionnaireAttempt).filter(QuestionnaireAttempt.user_id == user_id).delete()
    db.query(QuestionnaireSession).filter(QuestionnaireSession.user_id == user_id).delete()
    db.query(CounsellorProfile).filter(CounsellorProfile.user_id == user_id).delete()   # NEW
    db.query(User).filter(User.id == user_id).delete()
    db.query(UserAccess).filter(or_(UserAccess.user_id == user_id, UserAccess.counsellor_id == user_id)).delete()

    db.commit()
    return {"ok": True}

@app.post("/admin/emails/test") #test
def admin_test_email(to: EmailStr = Body(..., embed=True), admin=Depends(require_admin)):
    try:
        ok = send_email(
            to_email=str(to),
            subject="Psych Support: SMTP Test",
            html_body="<div style='font-family:system-ui'>✅ SMTP test from Psych Support.</div>",
            text_body="SMTP test from Psych Support (plain)."
        )
        return {"ok": ok}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"SMTP error: {e!r}")


@app.delete("/admin/users/{user_id}/chats")
def admin_clear_user_chats(
    user_id: int,
    purge_vectors: bool = False,
    db: Session = Depends(get_db),
    admin=Depends(require_admin),
):
    db.query(Message).filter(Message.user_id == user_id).delete()
    db.query(Conversation).filter(Conversation.user_id == user_id).delete()
    db.commit()
    if purge_vectors:
        delete_user_vectors(user_id)
    return {"ok": True}



@app.delete("/admin/users/{user_id}/questionnaires")
def admin_clear_user_questionnaires(user_id: int, db: Session = Depends(get_db), admin=Depends(require_admin)):
    db.query(QuestionnaireAttempt).filter(QuestionnaireAttempt.user_id == user_id).delete()
    db.query(QuestionnaireSession).filter(QuestionnaireSession.user_id == user_id).delete()
    db.commit()
    return {"ok": True}


@app.patch("/admin/users/{user_id}/password")
def admin_update_user_password(user_id: int, new_pw: str = Body(..., embed=True), db: Session = Depends(get_db), admin=Depends(require_admin)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.password_hash = hash_password(new_pw)
    db.commit()
    return {"ok": True}


@app.delete("/admin/vectorstore")
def admin_clear_vector_db(db: Session = Depends(get_db), admin=Depends(require_admin)):
    qc = init_qdrant()
    if qc:
        try:
            qc.delete_collection(COLLECTION)
        except Exception:
            # ignore if already deleted
            pass
    # also clear Document metadata table since they correspond to the vector store
    db.query(Document).delete()
    db.commit()
    return {"ok": True, "message": f"Vector collection '{COLLECTION}' deleted and document records cleared."}

@app.get("/admin/counsellors")
def admin_list_counsellors(db: Session = Depends(get_db), admin=Depends(require_admin)):
    rows = (
        db.query(User, CounsellorProfile)
        .join(CounsellorProfile, CounsellorProfile.user_id == User.id)
        .filter(User.is_counsellor == 1)
        .all()
    )
    out = []
    for u, c in rows:
        out.append({
            "id": u.id,
            "email": u.email,
            "name": c.name,
            "description": c.description,
            "experience_years": c.experience_years,
            "created_at": u.created_at.isoformat(),
        })
    return out


@app.post("/admin/counsellors/add")
def admin_add_counsellor(body: CounsellorCreate, db: Session = Depends(get_db), admin=Depends(require_admin)):
    user = db.query(User).filter(User.id == body.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # mark as counsellor
    user.is_counsellor = 1
    profile = CounsellorProfile(
        user_id=user.id,
        name=body.name.strip(),
        description=(body.description or "").strip(),
        experience_years=body.experience_years
    )
    db.merge(profile)
    db.commit()
    return {"ok": True, "user_id": user.id, "is_counsellor": True}

@app.get("/admin/users/non_counsellors")
def admin_list_non_counsellors(db: Session = Depends(get_db), admin=Depends(require_admin)):
    users = db.query(User).filter(User.is_counsellor == 0).all()
    return [{"id": u.id, "email": u.email} for u in users]

@app.patch("/admin/counsellors/{user_id}")
def admin_update_counsellor(user_id: int, body: CounsellorUpdate, db: Session = Depends(get_db), admin=Depends(require_admin)):
    c = db.query(CounsellorProfile).filter(CounsellorProfile.user_id == user_id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Counsellor not found")
    if body.name is not None: c.name = body.name
    if body.description is not None: c.description = body.description
    if body.experience_years is not None: c.experience_years = body.experience_years
    db.commit(); db.refresh(c)
    return {"ok": True}

@app.post("/admin/counsellors/{user_id}/demote")
def admin_demote_counsellor(user_id: int, db: Session = Depends(get_db), admin=Depends(require_admin)):
    u = db.query(User).filter(User.id == user_id).first()
    if not u or not u.is_counsellor:
        raise HTTPException(status_code=404, detail="Counsellor not found")
    u.is_counsellor = 0
    db.query(CounsellorProfile).filter(CounsellorProfile.user_id == user_id).delete()
    db.commit()
    return {"ok": True}

def _str_to_tags(s: Optional[str]) -> List[str]:
    if not s:
        return []
    return [t.strip() for t in s.split(",") if t.strip()]

@app.post("/counsellor/videos", response_model=VideoOut)
def counsellor_add_video(
    body: VideoCreate,
    db: Session = Depends(get_db),
    user=Depends(get_current_user)
):
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellor only")

    url = normalize_youtube_url(str(body.url))
    access = "all" if body.target_user_id is None else f"user:{body.target_user_id}"

    v = Video(
        id=str(uuid.uuid4()),
        title=body.title.strip(),
        url=url,
        platform="youtube",
        description=(body.description or "").strip() or None,
        tags=_tags_to_str(body.tags),   # store as string
        is_public=1 if body.is_public else 0,
        added_by=user.email,
        access=access,
    )
    # print(v)
    db.add(v); db.commit(); db.refresh(v)
    print("video added:", v.id, v.title, v.added_by, v.access)
    print(type(v.access))

    # convert for response
    return {
        **v.__dict__,
        "tags": _str_to_tags(v.tags)
    }
    
@app.delete("/counsellor/videos/{video_id}")
def counsellor_delete_video(video_id: str, db: Session = Depends(get_db), user=Depends(get_current_user)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admins should use /admin/videos delete")

    v = db.query(Video).filter(Video.id == video_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="Video not found")

    # ✅ Only allow deletion if this counsellor uploaded it
    if getattr(user, "is_counsellor", 0) != 1:
        raise HTTPException(status_code=403, detail="Only counsellors can delete videos")

    db.delete(v)
    db.commit() 
    return {"ok": True}

@app.get("/counsellor/users")
def counsellor_list_users(user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not getattr(user, "is_counsellor", False):
        raise HTTPException(status_code=403, detail="Counsellors only")
    users = db.query(User).filter(User.is_counsellor == 0).all()
    return [{"id": u.id, "email": u.email} for u in users]

@app.post("/user/access")
def grant_access(body: AccessCreate, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if is_admin_principal(user):
        raise HTTPException(status_code=403, detail="Admins cannot grant access.")
    # Verify counselor exists
    c = db.query(User).filter(User.id == body.counsellor_id, User.is_counsellor == 1).first()
    if not c:
        raise HTTPException(status_code=404, detail="Counsellor not found")

    # If an existing grant already exists for this pair, update it instead
    existing = db.query(UserAccess).filter(UserAccess.user_id == user.id, UserAccess.counsellor_id == body.counsellor_id).first()
    if existing:
        existing.allow_chats = 1 if body.allow_chats else 0
        existing.allow_q_chats = 1 if body.allow_q_chats else 0
        existing.allow_dashboard = 1 if body.allow_dashboard else 0
        db.commit()
        db.refresh(existing)

        # Send email notification for updated access
        send_email(
            to_email=c.email,
            subject="Access Updated",
            html_body=f"""
            <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
                <h2>Access Updated</h2>
                <p>User <b>{user.email}</b> has updated your access permissions:</p>
                <ul>
                    <li>Chats: {"Granted" if body.allow_chats else "Revoked"}</li>
                    <li>Dashboard: {"Granted" if body.allow_dashboard else "Revoked"}</li>
                </ul>
            </div>
            """
        )
        return {"ok": True, "id": existing.id, "updated": True}

    # Create a new access record
    acc = UserAccess(
        user_id=user.id,
        counsellor_id=body.counsellor_id,
        allow_chats=1 if body.allow_chats else 0,
        allow_q_chats=1 if body.allow_q_chats else 0,
        allow_dashboard=1 if body.allow_dashboard else 0,
    )
    db.add(acc)
    db.commit()
    db.refresh(acc)

    # Send email notification for granted access
    send_email(
        to_email=c.email,
        subject="Access Granted",
        html_body=f"""
        <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
            <h2>Access Granted</h2>
            <p>User <b>{user.email}</b> has granted you access to the following:</p>
            <ul>
                <li>Chats: {"Granted" if body.allow_chats else "Not Granted"}</li>
                <li>Dashboard: {"Granted" if body.allow_dashboard else "Not Granted"}</li>
            </ul>
        </div>
        """
    )
    return {"ok": True, "id": acc.id}


@app.get("/user/access")
def list_my_access(user=Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(UserAccess, User).join(User, User.id == UserAccess.counsellor_id).filter(UserAccess.user_id == user.id).all()
    out = []
    for acc, c in rows:
        out.append({
            "id": acc.id,
            "counsellor_id": acc.counsellor_id,
            "counsellor_email": c.email,
            "allow_chats": bool(acc.allow_chats),
            "allow_q_chats": bool(acc.allow_q_chats),
            "allow_dashboard": bool(acc.allow_dashboard),
            "created_at": acc.created_at.isoformat(),
        })
    return out


@app.patch("/user/access/{access_id}")
def update_my_access(access_id: int, body: AccessUpdate, user=Depends(get_current_user), db: Session = Depends(get_db)):
    acc = db.query(UserAccess).filter(UserAccess.id == access_id, UserAccess.user_id == user.id).first()
    if not acc:
        raise HTTPException(status_code=404, detail="Access record not found")
    if body.allow_chats is not None: acc.allow_chats = 1 if body.allow_chats else 0
    if body.allow_q_chats is not None: acc.allow_q_chats = 1 if body.allow_q_chats else 0
    if body.allow_dashboard is not None: acc.allow_dashboard = 1 if body.allow_dashboard else 0
    db.commit(); db.refresh(acc)
    return {"ok": True}


@app.delete("/user/access/{access_id}")
def delete_my_access(access_id: int, user=Depends(get_current_user), db: Session = Depends(get_db)):
    acc = db.query(UserAccess).filter(UserAccess.id == access_id, UserAccess.user_id == user.id).first()
    if not acc:
        raise HTTPException(status_code=404, detail="Access record not found")

    # Get the counselor's email before deleting the record
    c = db.query(User).filter(User.id == acc.counsellor_id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Counsellor not found")

    # Delete the access record
    db.delete(acc)
    db.commit()

    # Send email notification for revoked access
    send_email(
        to_email=c.email,
        subject="Access Revoked",
        html_body=f"""
        <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
            <h2>Access Revoked</h2>
            <p>User <b>{user.email}</b> has revoked your access permissions.</p>
        </div>
        """
    )
    return {"ok": True}

# -------- Counsellor: view grants given to them --------
@app.get("/counsellor/access")
def counsellor_access(user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellors only")
    rows = db.query(UserAccess, User).join(User, User.id == UserAccess.user_id).filter(UserAccess.counsellor_id == user.id).all()
    out = []
    for acc, u in rows:
        out.append({
            "id": acc.id,
            "user_id": u.id,
            "user_email": u.email,
            "allow_chats": bool(acc.allow_chats),
            "allow_q_chats": bool(acc.allow_q_chats),
            "allow_dashboard": bool(acc.allow_dashboard),
            "created_at": acc.created_at.isoformat(),
        })
    return out

# -------- Counsellor: fetch sessions for a user (only when allowed) --------
@app.get("/counsellor/users/{user_id}/sessions")
def counsellor_user_sessions(user_id: int, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellors only")
    acc = db.query(UserAccess).filter(UserAccess.user_id == user_id, UserAccess.counsellor_id == user.id, UserAccess.allow_chats == 1).first()
    if not acc:
        raise HTTPException(status_code=403, detail="Access to chats not granted by user")
    rows = db.query(Conversation).filter(Conversation.user_id == user_id).order_by(Conversation.created_at.desc()).all()
    return [{"id": r.id, "title": r.title, "created_at": r.created_at.isoformat()} for r in rows]


# -------- Counsellor: fetch messages in a user's session (only when allowed) --------
@app.get("/counsellor/users/{user_id}/sessions/{session_id}/messages")
def counsellor_user_session_messages(user_id: int, session_id: str, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellors only")
    acc = db.query(UserAccess).filter(UserAccess.user_id == user_id, UserAccess.counsellor_id == user.id, UserAccess.allow_chats == 1).first()
    if not acc:
        raise HTTPException(status_code=403, detail="Access to chats not granted by user")
    msgs = (
        db.query(Message)
        .filter(Message.conversation_id == session_id, Message.user_id == user_id)
        .order_by(Message.created_at.asc())
        .all()
    )
    return [{"role": m.role, "content": m.content, "created_at": m.created_at.isoformat()} for m in msgs]


# -------- Counsellor: fetch questionnaire dashboard for a user (only when allowed) --------
@app.get("/counsellor/users/{user_id}/questionnaire/dashboard")
def counsellor_user_questionnaire_dashboard(user_id: int, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellors only")
    acc = db.query(UserAccess).filter(UserAccess.user_id == user_id, UserAccess.counsellor_id == user.id, UserAccess.allow_dashboard == 1).first()
    if not acc:
        raise HTTPException(status_code=403, detail="Access to dashboard not granted by user")

    # reuse logic from /questionnaire/dashboard but for another user
    agg_rows = (
        db.query(
            QuestionnaireAttempt.day.label("day"),
            func.count(QuestionnaireAttempt.id).label("attempts"),
            func.avg(QuestionnaireAttempt.score).label("avg_score"),
        )
        .filter(QuestionnaireAttempt.user_id == user_id)
        .group_by(QuestionnaireAttempt.day)
        .order_by(desc(QuestionnaireAttempt.day))
        .all()
    )

    sub_latest = (
        db.query(
            QuestionnaireAttempt.day.label("day"),
            func.max(QuestionnaireAttempt.created_at).label("max_ts"),
        )
        .filter(QuestionnaireAttempt.user_id == user_id)
        .group_by(QuestionnaireAttempt.day)
        .subquery()
    )

    latest_rows = (
        db.query(QuestionnaireAttempt.day, QuestionnaireAttempt.comment)
        .join(sub_latest, (QuestionnaireAttempt.day == sub_latest.c.day) & (QuestionnaireAttempt.created_at == sub_latest.c.max_ts))
        .filter(QuestionnaireAttempt.user_id == user_id)
        .all()
    )
    latest_map = {d: c for d, c in latest_rows}

    out = []
    for r in agg_rows:
        out.append({
            "day": r.day,
            "attempts": int(r.attempts or 0),
            "avg_score": round(float(r.avg_score or 0.0), 2),
            "latest_comment": latest_map.get(r.day, None),
        })
    return out

@app.get("/public/counsellors")
def public_list_counsellors(db: Session = Depends(get_db)):
    rows = (
        db.query(User, CounsellorProfile)
        .join(CounsellorProfile, CounsellorProfile.user_id == User.id)
        .filter(User.is_counsellor == 1)
        .all()
    )
    out = [{"id": u.id, "email": u.email, "name": c.name, "experience_years": c.experience_years} for u,c in rows]
    return out

@app.post("/counsellor/request-access")
def request_access(
    payload: AccessRequestIn,
    db: Session = Depends(get_db),
    user=Depends(get_current_user)
):
    # Only counsellors allowed
    if not getattr(user, "is_counsellor", 0):
        raise HTTPException(status_code=403, detail="Counsellors only")

    # Check target user exists
    target = db.query(User).filter(User.id == payload.user_id).first()
    if not target:
        raise HTTPException(status_code=404, detail="Target user not found")

    # 🔥 Removed duplicate check — every call creates a new request
    req = AccessRequest(
        user_id=payload.user_id,
        counsellor_id=user.id,
        request_type=payload.request_type
    )
    db.add(req)
    db.commit()
    db.refresh(req)

    # ---- Send email ----
    send_email(
        target.email,
        "Access Request",
        f"Counsellor {user.email} is requesting access to your {payload.request_type}. "
        "Please log in to grant this request."
    )

    return {"ok": True, "msg": "Request sent", "request_id": req.id}



@app.post("/appointments")
def create_appointment(ap: AppointmentCreate, user=Depends(get_current_user), db: Session = Depends(get_db)):
    if user.is_counsellor:
        raise HTTPException(403, "Counsellor cannot create user appointment request")

    start = ap.start_time
    if start <= datetime.utcnow():
        raise HTTPException(400, "Cannot book an appointment in the past")

    end = start + timedelta(minutes=ap.duration_minutes)

    # Check overlap for both user and counsellor
    overlap = db.query(Appointment).filter(
        Appointment.status.in_(["pending","accepted"]),
        or_(
            and_(Appointment.user_id == user.id, Appointment.start_time < end, Appointment.end_time > start),
            and_(Appointment.counsellor_id == ap.counsellor_id, Appointment.start_time < end, Appointment.end_time > start)
        )
    ).first()
    if overlap:
        raise HTTPException(400, "Overlapping appointment exists")

    new_ap = Appointment(
        user_id=user.id,
        counsellor_id=ap.counsellor_id,
        start_time=start,
        end_time=end,
        duration_minutes=ap.duration_minutes,
        status="pending"
    )
    db.add(new_ap)
    db.commit()
    db.refresh(new_ap)

    # ✅ Email to counsellor
    counsellor = db.query(User).filter(User.id==ap.counsellor_id).first()
    if counsellor:
        send_email(
            counsellor.email,
            "New Appointment Request",
            f"Hello,\n\nYou have a new appointment request from {user.email}.\n\n"
            f"Date & Time: {start}\nDuration: {ap.duration_minutes} minutes\n\n"
            "Please log in to review."
        )

    return {"message": "Appointment requested", "appointment_id": new_ap.id}

@app.delete("/appointments/{ap_id}")
def delete_appointment(ap_id: int, user=Depends(get_current_user), db: Session = Depends(get_db)):
    ap = db.query(Appointment).get(ap_id)
    if not ap:
        raise HTTPException(404, "Not found")
    if ap.user_id != user.id and ap.counsellor_id != user.id:
        raise HTTPException(403, "Not allowed")

    u = db.query(User).get(ap.user_id)
    c = db.query(User).get(ap.counsellor_id)

    db.delete(ap)
    db.commit()

    # ✅ Send deletion emails
    if u:
        send_email(u.email, "Appointment Cancelled", f"Your appointment on {ap.start_time} was cancelled.")
    if c:
        send_email(c.email, "Appointment Cancelled", f"Appointment with {u.email if u else 'user'} on {ap.start_time} was cancelled.")

    return {"message": "Deleted"}

@app.get("/appointments/my")
def my_appointments(user=Depends(get_current_user), db: Session = Depends(get_db)):
    q = db.query(Appointment).filter(Appointment.user_id == user.id).order_by(Appointment.start_time.desc()).all()
    return [
        {
            "id": ap.id,
            "user_id": ap.user_id,
            "counsellor_id": ap.counsellor_id,
            "start_time": ap.start_time.isoformat(),
            "end_time": ap.end_time.isoformat(),
            "duration_minutes": ap.duration_minutes,
            "status": ap.status,
            "report": ap.report,  # Include the report field
            "meeting_link": ap.meeting_link,
        }
        for ap in q
    ]

@app.get("/appointments/counsellor")
def counsellor_appointments(user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not user.is_counsellor:
        raise HTTPException(403, "Only counsellors can view this page")

    q = db.query(Appointment).filter(Appointment.counsellor_id == user.id).order_by(Appointment.start_time.asc()).all()
    result = []
    for ap in q:
        u = db.query(User).get(ap.user_id)
        result.append({
            "id": ap.id,
            "user_email": u.email if u else "Unknown",
            "counsellor_id": ap.counsellor_id,
            "start_time": ap.start_time.isoformat(),
            "end_time": ap.end_time.isoformat(),
            "duration_minutes": ap.duration_minutes,
            "status": ap.status,
            "meeting_link": ap.meeting_link,
            "report": ap.report,
        })
    return result

@app.post("/appointments/{ap_id}/decision")
def decide_appointment(ap_id: int, decision: AppointmentDecision,
                       user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not user.is_counsellor:
        raise HTTPException(403, "Only counsellors can act on appointments")

    ap = db.query(Appointment).get(ap_id)
    if not ap or ap.counsellor_id != user.id:
        raise HTTPException(404, "Appointment not found")

    ap.status = decision.status
    if decision.status == "accepted":
        ap.meeting_link = decision.meeting_link
    db.commit()
    db.refresh(ap)

    # Email user
    u = db.query(User).get(ap.user_id)
    if u:
        if decision.status == "accepted":
            send_email(u.email, "Appointment Accepted",
                       f"Your appointment on {ap.start_time} has been accepted.\nMeeting link: {ap.meeting_link}")
        else:
            send_email(u.email, "Appointment Rejected",
                       f"Your appointment on {ap.start_time} was rejected.")

    return {"message": f"Appointment {decision.status}"}

@app.post("/appointments/{ap_id}/close")
def close_appointment(ap_id: int, close: AppointmentClose,
                      user=Depends(get_current_user), db: Session = Depends(get_db)):
    if not user.is_counsellor:
        raise HTTPException(403, "Only counsellors can close appointments")

    ap = db.query(Appointment).get(ap_id)
    if not ap or ap.counsellor_id != user.id:
        raise HTTPException(404, "Appointment not found")

    ap.status = "closed"
    ap.report = close.report
    db.commit()

    # Email user with report info
    u = db.query(User).get(ap.user_id)
    if u:
        send_email(u.email, "Appointment Closed",
                   f"Your appointment on {ap.start_time} has been closed.\nReport: {ap.report}")

    return {"message": "Appointment closed"}

# filepath: f:\Interview\llm-challenge\backend\main.py
@app.post("/appointments/counsellor/create")
def counsellor_create_appointment(
    body: CounsellorAppointmentCreate,
    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if not user.is_counsellor:
        raise HTTPException(status_code=403, detail="Only counsellors can create appointments.")

    # Log the incoming payload
    # print("Payload received:", body.dict())

    # Validate user_id
    target_user = db.query(User).filter(User.id == body.user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    # Validate start_time
    if body.start_time <= datetime.utcnow():
        raise HTTPException(status_code=400, detail="Cannot create appointment in the past")

    # Validate duration
    if body.duration_minutes <= 0:
        raise HTTPException(status_code=400, detail="Duration must be greater than 0")

    # Check for overlapping appointments
    end_time = body.start_time + timedelta(minutes=body.duration_minutes)
    overlap = (
        db.query(Appointment)
        .filter(
            Appointment.start_time < end_time,
            Appointment.end_time > body.start_time,
            or_(
                Appointment.user_id == body.user_id,
                Appointment.counsellor_id == user.id,
            ),
            Appointment.status.in_(["pending", "accepted"])
        )
        .first()
    )
    if overlap:
        raise HTTPException(status_code=400, detail="Appointment overlaps with an existing one.")

    # Create appointment
    appt = Appointment(
        user_id=body.user_id,
        counsellor_id=user.id,
        start_time=body.start_time,
        end_time=end_time,
        duration_minutes=body.duration_minutes,
        status="accepted",
        meeting_link=body.meeting_link,
    )
    db.add(appt)
    db.commit()
    db.refresh(appt)

    # Send email notification
    try:
        send_email(
            to_email=target_user.email,
            subject="New Appointment Scheduled",
            html_body=f"""
            <div style="font-family:system-ui,Segoe UI,Arial,sans-serif">
                <p>Your counsellor <b>{user.email}</b> has scheduled an appointment with you.</p>
                <p><b>Date & Time:</b> {body.start_time.strftime('%Y-%m-%d %H:%M')}</p>
                <p><b>Meeting link:</b> <a href="{body.meeting_link}">{body.meeting_link}</a></p>
            </div>
            """
        )
    except Exception as e:
        print(f"[MAIL] Failed to send appointment email: {e}")

    return {"message": "Appointment created by counsellor", "id": appt.id}


@app.get("/counsellor/users")
def list_users_for_counsellor(
    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if not user.is_counsellor and not isinstance(user, AdminPrincipal):
        raise HTTPException(status_code=403, detail="Only counsellors or admins can access this.")

    # Removed invalid `is_admin` filter
    q = db.query(User).filter(User.is_counsellor == False).all()
    return [{"id": u.id, "email": u.email} for u in q]